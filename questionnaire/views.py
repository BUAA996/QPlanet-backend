from result.views import delete_result
from django.http import JsonResponse
from django.shortcuts import render
from questionnaire.models import *
from user.models import *
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Max, query_utils
from django.utils import timezone
from QPlanet.values import *
from QPlanet.settings import *
from question.views import *
from result.views import *
import qrcode
import json
from datetime import datetime, timedelta
import random
import string
from django import utils
from django.db.models import Q
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx2pdf import convert
from random import randint as rand

def datetime_to_str(time):
	return time.strftime('%Y-%m-%d %H:%M')

def str_to_datetime(str):
	return datetime.strptime(str, '%Y-%m-%d %H:%M')
	
@csrf_exempt
def get_total(request):
	total = Questionnaire.objects.all().aggregate(Max('id'))
	if total['id__max'] == None:
		total = 1
	else:
		total = int(total['id__max'])
	total += 300

	s_total = SubmitInfo.objects.all().aggregate(Max('id'))
	if s_total['id__max'] == None:
		s_total = 1
	else:
		s_total = int(s_total['id__max'])
	s_total += 300

	return JsonResponse({'result': ACCEPT, 'message':r'获取成功!', 'total':total, 'submit_total':s_total})

def check_close(q):
	#已关闭返回1
	info = Info.objects.get(id = q.id)
	if info.state != RELEASE:
		return 1
	if q.deadline == None or datetime.now() < q.deadline:
		return 0
	info.state = SAVED
	info.save()
	return 1

def check_could_submit(q):
	if check_close(q):
		return 0
	if q.quota == None or q.count < q.quota:
		return 1

def hash(id):
	q = Questionnaire.objects.get(id = id)
	q.hash = ''.join(random.sample(string.ascii_letters + string.digits, 7)) + str(q.type)
	q.save()
	return q.hash

def build_questionnaire(title, description, own, type, deadline, quota, duration, random_order, select_less_score, certification, show_number, questions):
	total = Questionnaire.objects.all().aggregate(Max('id'))
	if total['id__max'] == None:
		total = 1
	else:
		total = int(total['id__max']) + 1
	if duration != None:
		duration = int(duration)
	Info.objects.create(id = total, state = SAVED)
	questionnaire = Questionnaire.objects.create(
		id = total, title = title, description = description, own = own, type = type, 
		deadline = deadline, quota = quota, duration = duration, count = 0, hash = "", 
		random_order = random_order, select_less_score = select_less_score, 
		certification = certification, show_number = show_number)
	hash_val = hash(questionnaire.id)
	save_questions(questions, questionnaire.id)
	
	return questionnaire.id, hash_val

@csrf_exempt
def create(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})

		data_json = json.loads(request.body)
		username = request.session.get('user')

		if data_json.get('deadline', -1) == -1 or data_json['deadline'] == None:
			temp = None
		else:
			temp = str_to_datetime(data_json['deadline'])
		id, hash = build_questionnaire(title = data_json['title'],
					description = data_json['description'],
					own = username,
					type = int(data_json['type']),
					deadline = temp,
					quota = data_json.get('quota', None),
					duration = data_json.get('duration', None),
					random_order = data_json.get('random_order', False),
					select_less_score = data_json.get('select_less_score', False),
					certification = int(data_json.get('certification', 0)),
					show_number = data_json.get('show_number', True),
					questions = data_json.get('questions', None)
			)

		return JsonResponse({'result': ACCEPT, 'message': r'保存成功!', 'id': id, 'hash': hash})
	else:
		print('IP is', request.META.get('HTTP_X_REAL_IP'))

@csrf_exempt
def list(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
		
		username = request.session.get('user')
		l = [x for x in Questionnaire.objects.filter(own = username)]
		l.sort(key = lambda x: x.id)
		l.reverse()
		result = {'result': ACCEPT, 'message': r'获取成功!', 'questionnaires':[]}
		for x in l:
			check_close(x)
			info = Info.objects.get(id = x.id)
			d = {'id': x.id, 'title': x.title, 'description': x.description, 'type': x.type,
				'count': x.count, 'hash': x.hash, 'state': info.state, 'quota': x.quota, 
				'create_time': datetime_to_str(x.create_time), 'deadline': x.deadline,
				'upload_time': datetime_to_str(info.upload_time)
			}
			d['create_time_int'] = int(x.create_time.timestamp())
			d['upload_time_int'] = int(info.upload_time.timestamp())
			result['questionnaires'].append(d)
		return JsonResponse(result)

@csrf_exempt
def reset_hash(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
		
		data_json = json.loads(request.body)
		id = int(data_json['id'])
		new = hash(id)

		return JsonResponse({'result': ACCEPT, 'message':r'重置成功!', 'hash':new})

@csrf_exempt
def delete(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
		
		data_json = json.loads(request.body)
		id = int(data_json['id'])
		q = Info.objects.get(id = id)
		if q.state == DELETED:
			questionnaire = Questionnaire.objects.get(id = q.id)
			delete_questions(id)
			delete_result(id)
			questionnaire.delete()
			q.delete()
			return JsonResponse({'result': ACCEPT, 'message':r'已彻底删除该问卷!'})
		else:
			q.state = DELETED
			q.save()
			return JsonResponse({'result': ACCEPT, 'message':r'已将该问卷移动至回收站!'})

@csrf_exempt
def recover(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
		
		data_json = json.loads(request.body)
		id = int(data_json['id'])
		q = Info.objects.get(id = id)
		q.state = SAVED
		q.save()
		return JsonResponse({'result': ACCEPT, 'message':r'已恢复!'})

@csrf_exempt
def release(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		id = int(data_json['id'])

		info = Info.objects.get(id = id)
		info.state = RELEASE
		info.save()

		_hash = hash(id)
		url = IMG_URL + _hash
		pic = qrcode.make(url)
		with open("img/"+_hash +".png","wb") as f:
			pic.save(f)
		return JsonResponse({'result': ACCEPT, 'message': r'发布成功!', 'url': url, 'img': IMG_URL + 'img/' + _hash + '.png'})

@csrf_exempt
def get_qr(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		id = int(data_json['id'])
		q = Questionnaire.objects.get(id = id)
		return JsonResponse({'result': ACCEPT, 'message':r'获取成功!', 'img':IMG_URL + 'img/' + q.hash + '.png'})

@csrf_exempt
def close(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		id = int(data_json['id'])

		info = Info.objects.get(id = id)
		info.state = SAVED
		info.save()

		return JsonResponse({'result': ACCEPT, 'message': r'已停止发布!'})

@csrf_exempt
def search_questionnaires(request):
	if request.method == 'POST':
		if request.session.get('is_login') != True:
			return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
		
		username = request.session.get('user')
		data_json = json.loads(request.body)
		
		q = data_json['query']
		res_tmp = []

		if not q.isdigit(): # 非仅数字，查标题
			l = [x for x in Questionnaire.objects.filter(Q(own = username) & Q(title__icontains = q))]
		else: # 仅数字
			l = [x for x in Questionnaire.objects.filter(Q(own = username) & Q(title__icontains = q))]
			l.extend([x for x in Questionnaire.objects.filter(Q(own = username) & Q(id = int(q)))])
		l.sort(key = lambda x: x.id)
		for x in l:
			check_close(x)
			info = Info.objects.get(id = x.id)
			d = {'id': x.id, 'title': x.title, 'description': x.description, 'type': x.type,
				'count': x.count, 'hash': x.hash, 'state': info.state, 'quota': x.quota, 
				'create_time': datetime_to_str(x.create_time), 'deadline': x.deadline,
				'upload_time': datetime_to_str(info.upload_time)
			}
			d['create_time_int'] = int(x.create_time.timestamp())
			d['upload_time_int'] = int(x.upload_time.timestamp())
			res_tmp.append(d)
		return JsonResponse({'result': ACCEPT, 'message': res_tmp})

@csrf_exempt
def check_type(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		hash = data_json['hash']
		if Questionnaire.objects.filter(hash = hash).exists() == False:
			return JsonResponse({'result': ERROR, 'message': r'问卷不存在!'})
		else:
			q = Questionnaire.objects.get(hash = hash)
			if check_close(q) == 1:
				return JsonResponse({'result': ERROR, 'message':r'问卷已关闭!'})
			return JsonResponse({'result': ACCEPT, 'message': r'获取成功!', 'requirement': int(hash[-1])})

@csrf_exempt
def view(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		if data_json.get('qid', -1) != -1:
			qid = int(data_json['qid'])
			if request.session.get('is_login') != True:
				return JsonResponse({'result': ERROR, 'message': r'您还未登录!'})
			q = Questionnaire.objects.get(id = qid)
			if q.own != request.session.get('user'):
				return JsonResponse({'result': ERROR, 'message': r'您没有权限!'})
		else:
			hash = data_json['hash']
			if Questionnaire.objects.filter(hash=hash).exists() == False:
				return JsonResponse({'result': ERROR, 'message': r'问卷不存在!'})
			q = Questionnaire.objects.get(hash = hash)
		
		result = {'result': ACCEPT, 'message': r'获取成功!',
				'qid':q.id,
				'title':q.title, 
				'description':q.description,
				'type': q.type,
				'show_number': q.show_number,
				'deadline': datetime_to_str(q.deadline),
				'quota': q.quota,
				'random_order': q.random_order,
				'select_less_score': q.select_less_score,
				'certification': q.certification,
				'show_number': q.show_number,
				'state': info.state
			}
		result['questions'] = get_questions(q.id) 
		return JsonResponse(result)

@csrf_exempt
def fill(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		hash = data_json['hash']
		if Questionnaire.objects.filter(hash = hash).exists() == False:
			return JsonResponse({'result': ERROR, 'message': r'问卷不存在!'})
		q = Questionnaire.objects.get(hash = hash)
		info = Info.objects.get(id = q.id)
		check_close(q)
		if info.state != RELEASE:
			return JsonResponse({'result': ERROR, 'message': r'问卷未发布!'})
		# TODO more information
		# Vote
		# Sign

		result = {'result': ACCEPT, 'message': r'获取成功!',
				'qid':q.id,
				'title':q.title, 
				'description':q.description,
				'type': q.type,
				'show_number': q.show_number,
				'deadline': datetime_to_str(q.deadline),
				'quota': q.quota,
				'random_order': q.random_order,
				'select_less_score': q.select_less_score,
				'certification': q.certification,
				'show_number': q.show_number,
				'state': info.state
			}
		
		result['questions'] = get_questions(q.id)
		if q.random_order == True:
			a = result['questions']
			l = len(a)
			# TODO set random seed
			random.seed(1)
			for i in range(30):
				x = rand(0, l-1)
				y = rand(0, l-1)
				if a[x]['is_essential'] == True or a[y]['is_essential'] == True:
					continue
				a[x],a[y] = a[y],a[x]
			result['questions'] = a
		return JsonResponse(result)

@csrf_exempt
def modify_questionnaire(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		
		if data_json.get('modify_type', -1) == -1 or data_json.get('qid', -1) == -1 \
			or data_json.get('title', -1) == -1 or data_json.get('description', -1) == -1 \
			or data_json.get('deadline', -1) == -1 or data_json.get('duration', -1) == -1 \
			or data_json.get('random_order', -1) == -1 or data_json.get('certification', -1) == -1 \
			or data_json.get('show_number', -1) == -1 or data_json.get('questions', -1) == -1 \
			or data_json.get('type', -1) == -1:
			return JsonResponse({'result': ERROR, 'message': FORM_ERROR})

		modify_type = data_json['modify_type']
		qid = data_json['qid']
		q = Questionnaire.objects.get(id = qid)

		if (q.type == 0 and data_json['type'] == 0) or (q.type == 5 and data_json['type'] == 5) \
			or (q.type in [1, 2, 3, 4] and data_json['type'] in [1, 2, 3, 4]) \
			or (q.type in [6, 7, 8, 9] and data_json['type'] in [6, 7, 8, 9]):

			if data_json['deadline'] == None:
				temp = None
			else:
				temp = str_to_datetime(data_json['deadline'])
			q.deadline = temp
			q.type = data_json['type']
			q.title = data_json['title']
			q.description = data_json['description']
			q.duration = data_json['duration']
			q.random_order = data_json['random_order']
			q.certification = data_json['certification']
			q.show_number = data_json['show_number']
			q.state = SAVED
			q.save()
			questions = data_json['questions']

			# 方式一：保留答卷；不能加题不能删题不能转换题目类型，可以移动题目不能移动选项，非考试类型可以加选项
			if modify_type == RESERVE_RESULTS:
				update_questions(questions)
			# 方式二：删除所有答卷（题目删掉重写）
			elif modify_type == DELETE_RESULTS:
				delete_questions(qid)
				save_questions(questions, qid)
				q.count = 0
				q.save()
				delete_result(qid)
			else:
				return JsonResponse({'result': ERROR, 'message': FORM_ERROR})
			return JsonResponse({'result': ACCEPT})
		else:
			return JsonResponse({'result': ERROR, 'message': r'问卷类型不可修改！'})

def copy_questionnaire(qid, title, to_username):
	q = Questionnaire.objects.get(id = qid)
	questions = get_questions(qid, False)
	res = build_questionnaire(title, q.description, to_username, q.type, q.deadline, q.duration, q.random_order, q.select_less_score, q.certification, q.show_number, questions)
	return res

@csrf_exempt
def copy_questionnaire_to_self(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		#return JsonResponse({'result': ACCEPT})
		qid = int(data_json['qid'])
		title = data_json['title']
		username = request.session.get('user')
		res = copy_questionnaire(qid, title, username)
		result = {'result': ACCEPT, 'copy_id': res[0], 'copy_hash': res[1]}
		return JsonResponse(result)

@csrf_exempt
def download(request):
	if request.method == 'POST':
		data_json = json.loads(request.body)
		if data_json.get('id', -1) != -1:
			id = int(data_json['id'])
			q = Questionnaire.objects.get(id = id)
		else:
			hash = data_json['hash']
			q = Questionnaire.objects.get(hash = hash)
		problems = [x for x in Question.objects.filter(questionnaire_id = q.id)]
		name = str(q.id)

		document = Document()
		Head = document.add_heading("",level=1)
		run  = Head.add_run(q.title)
		run.font.name=u'Cambria'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
		run.font.color.rgb = RGBColor(0,0,0)
		
		paragraph = document.add_paragraph("\n")
		count = 1
		for x in problems:
			if x.type in [SINGLE_CHOICE, MULTIPLE_CHOICE]:
				s = str(count) + "."
				if x.type == SINGLE_CHOICE:
					s += r'(单选) '
				else:
					s += r'(多选) '
				paragraph = document.add_paragraph(s + x.content)
				options = string_to_list(x.option)
				for j in range(len(options)):
					s = chr(j + 65) + '. '
					s = s + options[j]
					paragraph = document.add_paragraph(s)
				paragraph = document.add_paragraph("\n")
			elif x.type in [COMPLETION, DESCRIPTION]:
				if x.type == COMPLETION:
					s = str(count) + r".(填空) "
				else:
					s = str(count) + r".(简答) "
				paragraph = document.add_paragraph(s + x.content)
				paragraph = document.add_paragraph(r'___________________')
				if x.type == DESCRIPTION:
					paragraph = document.add_paragraph(r'___________________')
					paragraph = document.add_paragraph(r'___________________')
				paragraph = document.add_paragraph("\n")
			elif x.type == GRADING:
				s = str(count) + r".(打分) "
				# TODO add more detail
				paragraph = document.add_paragraph(s + x.content)
				paragraph = document.add_paragraph(r'___________________')
			else:
				s = str(count) + r".(定位) "
				paragraph = document.add_paragraph(s + x.content)
				paragraph = document.add_paragraph(r'___________________')
			count += 1
		
		document.save('img/' + name + '.docx')

		return JsonResponse({'result': ACCEPT, 'message':r'获取成功!', 
							'doc_name':name + '.docx', 'pdf_name':name + '.pdf'})