from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

document = Document()
Head = document.add_heading("",level=1)
run  = Head.add_run("问卷调查中文测试发电机房大大三大四大厦")
run.font.name=u'Cambria'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
run.font.color.rgb = RGBColor(0,0,0)

paragraph = document.add_paragraph(r'1.时间是?')
paragraph = document.add_paragraph(r'A. 11')
paragraph = document.add_paragraph(r'B. 22')

paragraph = document.add_paragraph("\n")

paragraph = document.add_paragraph(r'2. 填空')
paragraph = document.add_paragraph(r'___________________')

document.save('a.docx')